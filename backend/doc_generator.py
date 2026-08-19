# -*- coding: utf-8 -*-
"""用 python-docx 把抽取字段填入《半成品检验作业指导书》模版，生成 .docx。

跨平台（不依赖 MS Word / pywin32，纯 Python 操作 OOXML），取代旧版仅
Windows + 装了 Word 才能用的 win32com 方案。

模版 模版文件.docx 字段位置（对模版实际表格结构逐格探查所得，均为 0 基索引）：
  页眉表(2x6)：(0,1)=文件编号 (0,3)=过程名称 (0,5)=版本编号 (1,5)=发布日期
  正文表：      (0,0)=工序流程（整行合并；"工序检验" 4 字标灰底高亮）
               (2,1)=产品图号 (3,1)=产品名称 (4,1)=产品规格 (5,1)=材料
               (6,1)=产品颜色 (7,1)=产品净重 (8,1)=顾客代码
               (2,3)=产品图片+尺寸简图（横跨 5 列、纵跨 row2~8 的合并区）
  检验项目表从第 9 行（表头）开始，第 10 行起按类别分组，col0=项目（类别名，
  纵向合并），col1=序号，col2=检验项目（横跨 2 列），col4=特性标识，
  col5=检验设施/器具，col6=检验频次，col7=备注（跨类别合并注释，不动）。
  类别分组按 col0 的 w:vMerge 语义动态识别，不硬编码行号——这样「几何尺寸」
  「外观」等类别插行/清空多余行后，后续类别位置仍能正确定位。
"""

import copy
import os
import threading
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table, _Cell

import config

_DATA_START_ROW = 10  # 检验项目表数据行起始行号（第 9 行是表头）

# 统一 JSON 中文 key → 页眉表 (row, col)
_HEADER_MAP = {
    "文件编号": (0, 1), "过程名称": (0, 3), "版本编号": (0, 5),
    "发布日期": (1, 5),
}
# 统一 JSON 中文 key → 正文表 (row, col)
_BODY_MAP = {
    "产品图号": (2, 1), "产品名称": (3, 1), "产品规格": (4, 1),
    "材料": (5, 1), "产品颜色": (6, 1), "产品净重": (7, 1),
    "顾客代码": (8, 1),
}
_FLOW_CELL = (0, 0)          # 工序流程（整行合并）
_IMAGE_CELL = (2, 3)         # 产品图片+尺寸简图（合并区左上角）
_HIGHLIGHT_TARGET = "工序检验"
_HIGHLIGHT_COLOR = WD_COLOR_INDEX.GRAY_25

# 页脚表「编 制\n（日期）」值格 (row, col)：与正文表是两张独立的表，
# 位于 doc.sections[0].footer.tables[0]
_FOOTER_AUTHOR_CELL = (3, 3)

# 检验项目行字段 → 列号
_INSPECTION_COLMAP = [
    ("序号", 1), ("检验项目", 2), ("特性标识", 4),
    ("检验设施/器具", 5), ("检验频次", 6),
]

_gen_lock = threading.Lock()  # 输出文件按图纸名可能重名，简单串行避免并发写同一文件


def _row_col0_tc(table: Table, row_idx: int):
    """取某物理行 col0 的原始 <w:tc> 元素。

    不能用 table.cell(row_idx, 0)：python-docx 对纵向合并（w:vMerge）的续格会
    自动重定向、返回合并区首格的 _Cell，导致每一行都被误判成"新的一格"、读不到
    真实的 vMerge 状态。这里直接从该行自身的 <w:tr> 取第 0 个 <w:tc>（col0 在
    本模版里从不横向合并，tc_lst[0] 就是 col0），拿到未被重定向的真实单元格。
    """
    return table.rows[row_idx]._tr.tc_lst[0]


def _v_merge_tc(tc) -> Optional[str]:
    """读取（未被合并重定向的）原始 <w:tc> 的 w:vMerge 状态：
    None(未合并) / "restart"（合并首格）/ "continue"（合并续格，OOXML 里
    体现为存在 <w:vMerge> 但无 w:val 属性）。"""
    tcPr = tc.tcPr
    if tcPr is None:
        return None
    vm = tcPr.find(qn("w:vMerge"))
    if vm is None:
        return None
    return vm.get(qn("w:val")) or "continue"


def _set_cell_text(cell: _Cell, text: str) -> None:
    """设置单元格文本（支持多行），尽量沿用原有首个 run 的字体/字号/加粗，
    避免看起来和模版其余文字风格不一致。"""
    p = cell.paragraphs[0]
    src_run = p.runs[0] if p.runs else None
    font_name = src_run.font.name if src_run else None
    font_size = src_run.font.size if src_run else None
    bold = src_run.font.bold if src_run else None

    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)
    for run in list(p.runs):
        run._r.getparent().remove(run._r)

    lines = str(text).split("\n")
    run = p.add_run(lines[0])
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size
    if bold is not None:
        run.font.bold = bold


def _set_flow_with_highlight(cell: _Cell, flow: str) -> None:
    """写入「工序流程：xxx」，若 xxx 中含「工序检验」四字，对该子串加灰底高亮。"""
    text = "工序流程：" + flow
    p = cell.paragraphs[0]
    src_run = p.runs[0] if p.runs else None
    font_name = src_run.font.name if src_run else None
    font_size = src_run.font.size if src_run else None

    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)
    for run in list(p.runs):
        run._r.getparent().remove(run._r)

    idx = text.find(_HIGHLIGHT_TARGET)
    parts: List[Tuple[str, bool]] = (
        [(text, False)] if idx < 0 else
        [(text[:idx], False), (_HIGHLIGHT_TARGET, True),
         (text[idx + len(_HIGHLIGHT_TARGET):], False)]
    )
    for seg, hl in parts:
        if not seg:
            continue
        run = p.add_run(seg)
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if hl:
            run.font.highlight_color = _HIGHLIGHT_COLOR


def _iter_category_rows(table: Table) -> List[Tuple[str, int, int]]:
    """把检验项目表数据行按 col0 的 vMerge 语义分组，返回
    [(类别名, 起始行号, 现有行数), ...]（按出现顺序，0 基行号）。"""
    groups: List[Tuple[str, int, int]] = []
    cur_name: Optional[str] = None
    cur_start = 0
    n_rows = len(table.rows)
    for r in range(_DATA_START_ROW, n_rows):
        tc = _row_col0_tc(table, r)
        if _v_merge_tc(tc) == "continue" and cur_name is not None:
            continue
        if cur_name is not None:
            groups.append((cur_name, cur_start, r - cur_start))
        cur_name = _Cell(tc, table).text.strip()
        cur_start = r
    if cur_name is not None:
        groups.append((cur_name, cur_start, n_rows - cur_start))
    return groups


def _clear_row_data(table: Table, row: int) -> None:
    """清空多余行的全部数据列（含序号），避免留下和模版默认序号对不上的残影。"""
    for _key, col in _INSPECTION_COLMAP:
        _set_cell_text(table.cell(row, col), "")


def _insert_rows_after(table: Table, template_row: int, count: int) -> None:
    """在 template_row（该类别最后一行）之后复制插入 count 行（深拷贝该行
    的 <w:tr>，天然继承其各列的 gridSpan/vMerge 结构，包括 col0 的纵向合并
    延续、以及个别列自身可能存在的纵向合并——如「性能」类别的检验频次列）。"""
    anchor_tr = table.rows[template_row]._tr
    for _ in range(count):
        new_tr = copy.deepcopy(anchor_tr)
        anchor_tr.addnext(new_tr)
        anchor_tr = new_tr


def _fill_category(table: Table, cat_name: str, rows: list) -> None:
    """填某检验项目类别的明细行：现有行不够则复制末行插入，多了则清空多余行。"""
    if not rows:
        return
    groups = {name: (start, cnt) for name, start, cnt in _iter_category_rows(table)}
    if cat_name not in groups:
        return
    start, have = groups[cat_name]
    need = len(rows)

    if need > have:
        _insert_rows_after(table, start + have - 1, need - have)
    elif need < have:
        for r in range(start + need, start + have):
            _clear_row_data(table, r)

    for i, row in enumerate(rows):
        r = start + i
        for key, col in _INSPECTION_COLMAP:
            v = row.get(key, "")
            _set_cell_text(table.cell(r, col), "" if v is None else str(v))


def _insert_images(cell: _Cell, image_paths: List[str]) -> None:
    """把图片以「三列多行」插入图片区：清空占位文字，内嵌一个 3 列表格，
    每格一张图，按固定宽度缩放（锁定纵横比）。"""
    paths = [p for p in image_paths if p and os.path.isfile(p)]
    if not paths:
        return
    # 清空占位文字“产品图片+尺寸简图”
    for extra_p in cell.paragraphs[1:]:
        extra_p._p.getparent().remove(extra_p._p)
    for run in list(cell.paragraphs[0].runs):
        run._r.getparent().remove(run._r)

    rows = (len(paths) + 2) // 3
    nested = cell.add_table(rows, 3)
    try:
        nested.style = None
    except Exception:  # noqa: BLE001
        pass
    target_w = Inches(1.6)
    for idx, p in enumerate(paths):
        r, c = divmod(idx, 3)
        run = nested.cell(r, c).paragraphs[0].add_run()
        run.add_picture(os.path.abspath(p), width=target_w)


def generate(diagram_name: str, data: dict, image_relpaths=None) -> str:
    """把统一 JSON（中文 key + 检验项目嵌套数组）填入模版，可选把选中图片插入图片区，
    生成并返回 .docx 路径。image_relpaths 为相对 图纸/ 目录的 '/' 分隔路径列表。"""
    vals = data or {}
    image_paths = [
        os.path.join(config.DIAGRAMS_DIR, rel.replace("/", os.sep))
        for rel in (image_relpaths or [])
    ]
    stem = diagram_name[:-4] if diagram_name.lower().endswith(".pdf") else diagram_name
    out_dir = os.path.join(config.PROJECT_ROOT, "outputs")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"{stem}_工序检验作业指导书.docx")

    tpl = os.path.join(config.PROJECT_ROOT, "模版文件.docx")
    if not os.path.isfile(tpl):
        raise RuntimeError("未找到模版文件.docx")

    with _gen_lock:
        doc = Document(tpl)

        # 页眉字段
        htab = doc.sections[0].header.tables[0]
        for k, (r, c) in _HEADER_MAP.items():
            if vals.get(k):
                _set_cell_text(htab.cell(r, c), str(vals[k]))

        # 正文表：产品信息字段
        btab = doc.tables[0]
        for k, (r, c) in _BODY_MAP.items():
            if vals.get(k):
                _set_cell_text(btab.cell(r, c), str(vals[k]))

        # 页脚「编制（日期）」：姓名、日期分两行，日期不加括号
        author = str(vals.get("编制", "") or "").strip()
        author_date = str(vals.get("日期", "") or "").strip()
        if author or author_date:
            ftab = doc.sections[0].footer.tables[0]
            text = "\n".join(s for s in (author, author_date) if s)
            r, c = _FOOTER_AUTHOR_CELL
            _set_cell_text(ftab.cell(r, c), text)

        # 工序流程 + “工序检验”灰底高亮
        flow = str(vals.get("工序流程", "") or "")
        if flow:
            fr, fc = _FLOW_CELL
            _set_flow_with_highlight(btab.cell(fr, fc), flow)

        # 产品图片+尺寸简图（三列多行）
        ir, ic = _IMAGE_CELL
        _insert_images(btab.cell(ir, ic), image_paths)

        # 检验项目表：逐类别填（几何尺寸/外观可变行数会插行或清空多余行）
        for cat in vals.get("检验项目", []):
            _fill_category(btab, cat.get("项目", ""), cat.get("明细", []))

        doc.save(out_path)
        return out_path
